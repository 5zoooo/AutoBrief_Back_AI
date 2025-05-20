from docx import Document
import os

class WordExporterAgent:
    def __init__(self, output_path: str):
        self.output_path = output_path

    def export(self, structure: dict) -> str:
        os.makedirs(os.path.dirname(self.output_path), exist_ok=True)
        doc = Document()

        for block in structure.get("structure", []):
            if block["type"] == "title":
                doc.add_heading(block["content"], level=1)
            elif block["type"] == "subtitle":
                doc.add_heading(block["content"], level=2)
            elif block["type"] == "paragraph":
                doc.add_paragraph(block["content"])
            elif block["type"] == "list":
                for item in block["content"]:
                    doc.add_paragraph(item, style="List Bullet")
            elif block["type"] == "table":
                rows = len(block["content"])
                cols = len(block["content"][0])
                table = doc.add_table(rows=rows, cols=cols)
                for i, row in enumerate(block["content"]):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell

        doc.save(self.output_path)
        return self.output_path
